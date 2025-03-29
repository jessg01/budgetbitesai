import os
import docx
import difflib
import re
from ollama import chat
import sys
import time # Added for simulating confidence pause ;)

# --- Configuration ---
OLLAMA_MODEL = 'gemma3:4b'
MEALS_DOCX_PATH = 'meals.docx'
WALMART_DOCX_PATH = 'walmart.docx'
MAX_SELECTION = 5

# --- Helper Functions ---

def exit_with_error(message):
    """Prints an error message and exits."""
    print(f"\nERROR: {message}", file=sys.stderr)
    sys.exit(1)

def parse_docx_lines(file_path, error_noun="file"):
    """Parses non-empty, cleaned lines from a Word document."""
    if not os.path.exists(file_path):
        exit_with_error(f"{error_noun.capitalize()} not found at '{file_path}'")
    try:
        document = docx.Document(file_path)
        items = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                cleaned = re.sub(r"^[•*\-–—]?\s*", "", text).strip()
                if cleaned:
                    items.append(cleaned)
        if not items and error_noun != "file": # Add specific check for content files
             exit_with_error(f"No text content found in {error_noun} file: '{file_path}'")
        return items
    except Exception as e:
        exit_with_error(f"Could not read or parse {error_noun} '{file_path}': {e}")

def parse_meals(file_path):
    """Parses meal blocks from the meals document."""
    if not os.path.exists(file_path):
         exit_with_error(f"Meals file not found at '{file_path}'")
    try:
        document = docx.Document(file_path)
        paragraphs = document.paragraphs
        meal_blocks = []
        i = 0
        while i < len(paragraphs):
            current_text = paragraphs[i].text.strip()
            if "time:" in current_text.lower() and "shelf life:" in current_text.lower():
                meal_block = [current_text]
                i += 1
                while i < len(paragraphs):
                    next_text = paragraphs[i].text.strip()
                    if not next_text:
                        i += 1
                        continue
                    if "time:" in next_text.lower() and "shelf life:" in next_text.lower():
                        break
                    meal_block.append(next_text)
                    i += 1
                meal_blocks.append("\n".join(meal_block))
            else:
                i += 1
        if not meal_blocks:
             exit_with_error("No meal blocks found. Check format (needs lines with 'Time:' and 'Shelf life:').")
        return meal_blocks
    except Exception as e:
        exit_with_error(f"Error reading or parsing meals docx file '{file_path}': {e}")

def get_user_meal_selection(meals_list, max_select):
    """Gets user meal choices by number or name, returns list of (name, details) tuples."""
    selected = []
    processed_inputs = set()

    print(f"\nEnter up to {max_select} meal names or numbers (comma-separated):")
    user_input = input("> ").strip()
    inputs = [item.strip() for item in user_input.split(",") if item.strip()]

    if not inputs:
        exit_with_error("No meal selection provided.")

    meal_dict = {str(i+1): entry for i, entry in enumerate(meals_list)}
    name_dict = {entry[0].lower(): entry for entry in meals_list}
    all_names = [m[0] for m in meals_list]

    for item in inputs[:max_select]:
        if item in processed_inputs:
            continue

        meal_entry = None
        item_lower = item.lower()

        if item in meal_dict: meal_entry = meal_dict[item]
        elif item_lower in name_dict: meal_entry = name_dict[item_lower]
        else:
            close_matches = difflib.get_close_matches(item, all_names, n=1, cutoff=0.7)
            if close_matches:
                matched_name = close_matches[0]
                matched_name_lower = matched_name.lower()
                if matched_name_lower in name_dict:
                    temp_entry = name_dict[matched_name_lower]
                    is_new = all(temp_entry[0] != sel[0] for sel in selected)
                    if is_new:
                        print(f"(Interpreting '{item}' as '{matched_name}')")
                        meal_entry = temp_entry
                    else:
                         processed_inputs.add(item); continue

        if meal_entry:
             is_new = all(meal_entry[0] != sel[0] for sel in selected)
             if is_new:
                 selected.append(meal_entry)
                 processed_inputs.add(item)
                 processed_inputs.add(meal_entry[0].lower())
                 for num_str, entry_val in meal_dict.items():
                     if entry_val[0] == meal_entry[0]: processed_inputs.add(num_str); break
             else:
                 processed_inputs.add(item) # Mark input as processed (duplicate)
        else:
             print(f"Warning: Could not find meal matching '{item}'. Skipping.")
             processed_inputs.add(item)

    if not selected:
        exit_with_error("No valid meals were ultimately selected.")

    # Final deduplication
    unique_selected = []
    seen_names = set()
    for meal in selected:
        if meal[0] not in seen_names:
            unique_selected.append(meal)
            seen_names.add(meal[0])
    return unique_selected


def call_ollama_non_streaming(model, messages, task_description="LLM task"):
    """Calls Ollama non-streamingly, returns content or None on error."""
    try:
        response = chat(model=model, messages=messages, stream=False)
        if response and 'message' in response and 'content' in response['message']:
            return response['message']['content'].strip()
        else:
            print(f"\nError: Invalid response structure from Ollama for {task_description}.")
            print("Response received:", response, file=sys.stderr)
            return None
    except Exception as e:
        print(f"\nError: Ollama call failed for {task_description}: {e}", file=sys.stderr)
        print(f"Ensure Ollama server running & model '{model}' available.", file=sys.stderr)
        return None

def parse_and_sum_prices(text_output):
    """Parses the LLM comparison output and sums estimated prices."""
    total_cost = 0.0
    # Regex to find lines with the expected format and capture the price
    # Looks for "->" then some text, then "(Estimated Price: $NUMBER)"
    # Allows for variations in spacing around the price.
    price_pattern = re.compile(r'->.*\(Estimated Price:\s*\$([\d,]+\.?\d*)\)', re.IGNORECASE)

    lines = text_output.splitlines()
    for line in lines:
        match = price_pattern.search(line)
        if match:
            price_str = match.group(1).replace(',', '') # Remove commas for float conversion
            try:
                total_cost += float(price_str)
            except ValueError:
                print(f"Warning: Could not parse price '{match.group(1)}' in line: {line}", file=sys.stderr)
    return total_cost

# --- Main Execution ---

# 1. Load and Prepare Meals
meal_blocks = parse_meals(MEALS_DOCX_PATH)
meals_ordered = []
for block in meal_blocks:
    lines = block.splitlines()
    if lines:
        meal_name = re.sub(r"^[•*\-–—]?\s*", "", lines[0].strip())
        if meal_name:
            meals_ordered.append((meal_name, block))
if not meals_ordered: exit_with_error("Failed to extract valid meal names and details.")

# 2. Display Meals and Get User Selection
print("--- Available Meals ---")
for idx, (meal_name, _) in enumerate(meals_ordered, start=1): print(f"{idx}. {meal_name}")
print("-----------------------")
selected_meals = get_user_meal_selection(meals_ordered, MAX_SELECTION)

# 3. Consolidate Selected Meal Details
consolidated_info = "\n---\n".join([f"Meal: {name}\n{details}" for name, details in selected_meals])

# 4. Generate Initial Grocery List (Non-Streaming)
print("\nGenerating intermediate grocery list...")
system_prompt_grocery = ("Concise assistant: Create a consolidated grocery list from meal details, combining quantities. Only use mentioned ingredients.")
user_prompt_grocery = f"Meal Details:\n{consolidated_info}\n\nGenerate the consolidated grocery list."
messages_grocery = [{'role': 'system', 'content': system_prompt_grocery}, {'role': 'user', 'content': user_prompt_grocery}]
llm_grocery_list = call_ollama_non_streaming(OLLAMA_MODEL, messages_grocery, "Grocery List Generation")
if not llm_grocery_list: exit_with_error("Failed to generate the initial grocery list.")

# 5. Load Walmart Items
walmart_items = parse_docx_lines(WALMART_DOCX_PATH, error_noun="Walmart items")
walmart_list_str = "\n".join(sorted(walmart_items))

# 6. Define Comparison Prompt with Pricing Instruction
system_prompt_comparison = (
    "You are a precise shopping assistant. Compare the 'Generated Grocery List' to 'Available Walmart Items'. "
    "For each grocery item, find the closest Walmart match. Crucially, for every matched item, provide a confident estimated price in USD ($). "
    "Even if approximating, state the price confidently."
)
user_prompt_comparison = (
    f"Generated Grocery List:\n```\n{llm_grocery_list}\n```\n\n"
    f"Available Walmart Items:\n```\n{walmart_list_str}\n```\n\n"
    "Task: For each item in the 'Generated Grocery List':\n"
    "1. Find the *closest* match from 'Available Walmart Items'.\n"
    "2. State the match and provide a *confidently estimated price* for the matched Walmart item.\n"
    "Use this exact format for matches:\n"
    '"Grocery Item" -> "Walmart Product Match" (Estimated Price: $X.XX)\n'
    "If no close match is found, use this exact format:\n"
    '"Grocery Item" -> "No close match found in Walmart list"\n'
    "Be precise. Only estimate prices for items you successfully match to the Walmart list."
)
messages_comparison = [{'role': 'system', 'content': system_prompt_comparison}, {'role': 'user', 'content': user_prompt_comparison}]

# 7. Generate and Stream Walmart Comparison & Capture Output
print(f"\n--- Recommended Walmart Purchases (Streaming from {OLLAMA_MODEL}) ---")
full_comparison_output = "" # Initialize buffer to capture the full response
try:
    stream = chat(
        model=OLLAMA_MODEL,
        messages=messages_comparison,
        stream=True
    )
    for chunk in stream:
        if 'message' in chunk and 'content' in chunk['message']:
            content_piece = chunk['message']['content']
            print(content_piece, end='', flush=True) # Stream to console
            full_comparison_output += content_piece      # Append to buffer
    print() # Final newline for console output

except Exception as e:
    print(f"\nError during Ollama streaming call: {e}", file=sys.stderr)
    print(f"Ensure Ollama server running & model '{OLLAMA_MODEL}' available.", file=sys.stderr)
    # Decide if exit is needed - maybe allow cost calculation attempt if some output exists?
    if not full_comparison_output: # Exit if we got nothing useful
         exit_with_error("Streaming failed before any output was captured.")

print("-----------------------------------------------------------------")

# 8. Calculate and Display Total Estimated Cost
if full_comparison_output:
    print("\nCalculating estimated total cost...")
    total_estimated_cost = parse_and_sum_prices(full_comparison_output)
    time.sleep(1) # Simulate thoughtful calculation for confidence effect ;)
    print("\n===================================")
    print("  CONFIDENT COST ESTIMATION")
    print("-----------------------------------")
    print(f" The total estimated cost for the matched items is: ${total_estimated_cost:.2f}")
    print("===================================")
    if total_estimated_cost == 0:
        print("(Note: No prices could be extracted from the AI's response. Ensure the model followed the requested format.)")
else:
    print("\nCould not calculate total cost as no comparison output was captured.")